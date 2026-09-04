#!/usr/bin/env python3
"""Regenerates willed.docx to the DOCX binding's locked engineered form (WILL-1.md,
"DOCX -- hidden-text marker lines"): a marker is one paragraph containing
a single hidden run (w:vanish on the run), and the paragraph mark itself
is ALSO hidden (w:vanish in pPr/rPr) so the marker paragraph contributes
no vertical space and a willed document is layout-identical to its
unwilled twin.

Run from anywhere:  python3 generate_docx.py
Requires: python-docx (tested at 1.2.0)
Builds a fresh document -- no copied parts, no stray customXml.
"""
import os
import re
import zipfile

import docx
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DEFAULT_OUT = os.path.join(HERE, "willed.docx")


def marker_paragraph(doc, text):
    """One paragraph, one hidden run carrying the marker text verbatim,
    AND the paragraph mark's own run properties hidden too."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    rPr_mark = OxmlElement('w:rPr')
    rPr_mark.append(OxmlElement('w:vanish'))
    pPr.append(rPr_mark)

    run = p.add_run(text)
    run.font.hidden = True
    return p


def _strip_stray_customxml(out_path):
    """python-docx's built-in default.docx template carries an unused
    customXml part (a bibliography schema reference). No Will data lives
    there; strip it so the shipped example carries nothing stray."""
    with zipfile.ZipFile(out_path, "r") as zin:
        entries = [(i, zin.read(i.filename)) for i in zin.infolist()]

    def strip(name, data):
        if name.startswith("customXml/"):
            return None
        if name == "word/_rels/document.xml.rels":
            text = data.decode("utf-8")
            text = re.sub(r'<Relationship[^>]*Type="[^"]*relationships/customXml"[^>]*/>', "", text)
            return text.encode("utf-8")
        if name == "[Content_Types].xml":
            text = data.decode("utf-8")
            text = re.sub(r'<Override PartName="/customXml/[^"]*"[^>]*/>', "", text)
            return text.encode("utf-8")
        return data

    kept = []
    for info, data in entries:
        new_data = strip(info.filename, data)
        if new_data is None:
            continue
        # Zip entries carry a per-file modified timestamp; left alone it
        # is "now", which makes the archive bytes -- and so its hash --
        # differ between two otherwise-identical runs. Pin it so this
        # generator is actually reproducible, matching the fixed date
        # already baked into docProps/core.xml by the underlying
        # template.
        info.date_time = (2013, 12, 23, 23, 15, 0)
        kept.append((info, new_data))

    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for info, data in kept:
            zout.writestr(info, data)


def build(out_path=None):
    """Builds willed.docx and returns its path. Importable and callable
    directly -- unlike a plain top-to-bottom script, importing this
    module no longer regenerates the file as a side effect; callers
    (verify_docx.py included) call build() explicitly."""
    out_path = out_path or DEFAULT_OUT
    doc = docx.Document()

    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.title = ""
    cp.subject = ""
    cp.comments = ""

    doc.add_paragraph("Ordinary prose an agent may edit freely.")
    marker_paragraph(doc, "<!-- will/1 keep: the sentence legal approved; not ours to improve -->")
    doc.add_paragraph("The approved sentence.")
    marker_paragraph(doc, "<!-- /will -->")
    marker_paragraph(doc, "<!-- will/1 append: chronological entries only -->")
    doc.add_paragraph("- 2026-01-19 - Partial results are always reported.")
    marker_paragraph(doc, "<!-- /will -->")
    doc.add_paragraph("Closing prose, also free.")

    doc.save(out_path)
    _strip_stray_customxml(out_path)
    return out_path


if __name__ == "__main__":
    written = build()
    print("wrote", written, "(customXml stripped)")
